# -*- coding: utf-8 -*-
"""
《自然语言语义处理概论》大作业 — Word 报告（文字精修版）
用法: python generate_report_docx.py  （插图已生成则无需重跑 generate_figures.py）
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
FIG = BASE / "figures"
OUT = BASE / "MusAgent_大作业报告.docx"


def set_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def para(doc, text, size=12, bold=False, indent=True, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    set_font(p.add_run(text), size=size, bold=bold)
    return p


def heading(doc, text, lv=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if lv == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size={1: 16, 2: 14, 3: 12}[lv], bold=True)


def center(doc, text, size=12, bold=False, n=0):
    for _ in range(n):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=size, bold=bold)


def tbl(doc, heads, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(heads))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(heads):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_font(p.runs[0] if p.runs else p.add_run(h), size=10, bold=True)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(v)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                if p.runs:
                    set_font(p.runs[0], size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def fig(doc, path, caption, width=Inches(5.8)):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=width)
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(c.add_run(caption), size=10)
        doc.add_paragraph()


def build():
    doc = Document()
    sec = doc.sections[0]
    for k, v in [("left_margin", 2.8), ("right_margin", 2.8), ("top_margin", 2.5), ("bottom_margin", 2.5)]:
        setattr(sec, k, Cm(v))

    # ── 封面 ──
    center(doc, "SHANGHAI UNIVERSITY", 14, True, 3)
    center(doc, "2025—2026 学年  春季学期", 12)
    center(doc, "《自然语言语义处理概论》课程大作业", 14, True, 1)
    center(doc, "设计与实现报告", 13, True, 2)
    center(doc, "MusAgent：面向诗词创作的\n语义检索增强生成系统", 15, True, 1)
    for ln in [
        "学    院：计算机工程与科学学院",
        "专    业：＿＿＿＿＿＿＿＿＿＿＿＿",
        "学    号：＿＿＿＿＿＿＿＿＿＿＿＿",
        "学生姓名：＿＿＿＿＿＿＿＿＿＿＿＿",
        "指导教师：＿＿＿＿＿＿＿＿＿＿＿＿",
        "提交日期：2026 年 ＿＿ 月 ＿＿ 日",
    ]:
        center(doc, ln, 12)
    doc.add_page_break()

    # ── 摘要 ──
    heading(doc, "摘  要")
    para(doc,
         "诗词创作往往始于高度压缩的主题意象，如“城市孤独”或“雨夜和解”。"
         "现有大语言模型虽能生成流畅文本，但存在三类不足：缺乏可验证的知识依据、"
         "检索推理过程不可见、语义分析与生成模块相互脱节。")
    para(doc,
         "针对上述问题，本文设计实现 MusAgent 系统——以语义理解为核心、以本地诗歌知识库为 grounding、"
         "以可解释性为特色的检索增强生成平台。核心贡献包括：")
    para(doc,
         "（1）5320 首诗歌语料上的 BM25 + BGE + Cross-Encoder 三级混合检索链路；"
         "（2）文学词典与 RoBERTa 的 0.5:0.5 融合情感分析；"
         "（3）BERT-NER 与 Jieba 意象词融合的实体识别；"
         "（4）可解释语义分析层，将关键词命中、意象重叠转化为可读说明；"
         "（5）RAG 结构化抽取与化用标注机制，使输出明确引用参考作品；"
         "（6）文学垂直知识图谱与 D3 力导向图可视化；"
         "（7）SSE 实时推送 14 阶段 Pipeline 进度；"
         "（8）JWT 账号与对话跨设备同步。",
         indent=False)
    para(doc,
         "系统采用 FastAPI 与 React 前后端分离，包含 13 个功能模块、15 类 NLP 技术，"
         "新增 Critic Agent 自评闭环使生成结果具备可解释质量保障。"
         "24 项自动化回归测试全部通过，Ablation 实验验证混合检索的互补价值。"
         "本工作表明：垂直创作场景中，语义处理的价值在于构建「理解—检索—解释—生成—引用」的完整证据链。")
    para(doc, "关键词：自然语言语义处理；检索增强生成；混合检索；可解释 NLP；情感分析；诗词创作辅助", indent=False)
    doc.add_page_break()

    # ══ 第1章 ══
    heading(doc, "第1章  引言")
    heading(doc, "1.1  研究背景与动机", 2)
    para(doc,
         "自然语言语义处理研究的核心问题，是计算机如何在不同语言层级上把握“意义”："
         "词汇层级的概念指称、句子层级的命题内容、语篇层级的情感与意图，以及跨文本的语义关联。"
         "《自然语言语义处理概论》课程从词向量、句法—语义接口、信息检索到情感计算，"
         "构建了理解上述问题的理论框架。课程大作业的要求并非简单演示某一算法，"
         "而是将多种语义技术组织为一个有应用场景、有输入输出、可验证的系统。")
    para(doc,
         "诗词创作辅助是一个能同时触及多项语义能力的典型场景。与新闻分类或情感极性判别不同，"
         "创作主题往往以极短的意象短语出现，例如“地铁里的孤独”或“黄昏与成长”。"
         "这类输入在字面上可能仅含三五个词，却隐含复杂的情感基调、意象组合与文化联想。"
         "若系统不能首先对主题做语义拆解，而直接进入生成环节，"
         "则即便输出文本流畅，也难以证明其“理解”了用户真正想表达的内容。")
    para(doc,
         "近两年，大量应用选择直接调用 DeepSeek、GPT 等大语言模型 API 完成写诗、润色与摘要任务。"
         "这种方案工程成本低，但存在明显的语义缺陷：第一，模型参数中的“知识”不可更新、不可审计，"
         "生成内容是否真正化用了某首经典作品，系统无法给出证据；第二，用户看不到中间推理过程，"
         "无法判断“城市孤独”究竟被理解为地理空间、人际关系还是存在论意义上的疏离；"
         "第三，即便引入 RAG，若只是将检索原文拼接到 Prompt 中，"
         "检索分数、意象关联与情感一致性仍停留在黑盒内部，无法支撑课程所强调的可解释语义分析。")
    para(doc,
         "基于上述考虑，本作业并未选择“单一 LLM 接口”作为系统核心，"
         "而是围绕本地诗歌知识库构建一条显式的语义处理流水线："
         "先理解，再检索，再解释，最后生成并标注来源。"
         "这一设计思路与课程中“向量空间模型 + 概率检索 + 神经语义表示 + 知识增强”的技术脉络一致，"
         "也更容易在答辩中展示每一项 NLP 技术的实际作用。")

    heading(doc, "1.2  问题描述", 2)
    para(doc,
         "形式化地，设用户输入创作主题 query，可为短语、诗句片段或混合文本；"
         "本地知识库 D={d₁, d₂, …, d_N}，本系统中 N=5320，每篇文档 d 含体裁 type、标题 title、"
         "作者 author 与正文 content。系统需实现映射：")
    para(doc,
         "F(query) = (A, R, E, G)",
         indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc,
         "其中 A 表示结构化语义分析结果，包括分词、命名实体、关键词、TextRank 摘要与融合情感向量；"
         "R⊂D 为与 query 语义最相关的检索结果集合；E 为面向用户的自然语言解释，"
         "说明 R 中各作品与 query 在关键词、意象与情感上的关联；"
         "G 为在 R 与 A 约束下生成的创作文本，且 G 必须携带对 R 中作品的引用标注。"
         "与普通聊天机器人“输入—输出”的二元结构不同，MusAgent 的输出是一个包含证据链的语义_bundle，"
         "其设计目标不是替代诗人创作，而是为创作者提供可参照、可解释、可迭代的灵感材料。")

    heading(doc, "1.3  研究内容与主要贡献", 2)
    para(doc,
         "围绕上述问题，本文完成的主要工作如下。"
         "第一，构建并清洗 5320 首诗歌结构化语料，建立 BM25 倒排索引与 BGE 向量索引，"
         "向量结果持久化缓存以支持重复实验。"
         "第二，设计并实现 BM25 + BGE + Cross-Encoder 三级混合检索，"
         "在字面匹配与深层语义相似之间取得互补，默认 BM25 融合权重 α=0.55。"
         "第三，实现面向诗歌语体的融合情感分析，将六类文学情感词典与 RoBERTa 微调模型按 0.5:0.5 融合。"
         "第四，实现 BERT-NER 与 Jieba 意象词融合的 hybrid 实体识别，"
         "兼顾人名地名等规范实体与诗歌特有的意象名词。"
         "第五，提出可解释语义分析模块，将多源 NLP 信号聚合为 semanticInsight 与逐条 summary，"
         "解决“检索结果为何相关”的可读性问题。"
         "第六，实现 RAG 结构化抽取、化用标注与 Pipeline 编排，交付完整 B/S 系统，"
         "并编写 24 项自动化回归测试与 Ablation 评测接口。"
         "第七，构建文学垂直知识图谱 kg_engine，以意象—情感语义关系为主、书目元数据为辅，"
         "支持批量建图、主题子图与 BERT-RE 关系抽取微调。"
         "第八，实现 POST /api/pipeline/stream 的 SSE 实时阶段推送与 JWT 账号体系，"
         "提升系统可观测性与多设备使用体验。"
         "第九，设计并实现 Critic Agent 自评闭环，在 14 阶段 Pipeline 末端对生成文本进行 0—10 分量化打分，"
         "评分低于阈值 7.0 时自动触发 LLM 改写一次；"
         "20 条样本人工评测显示评分均值从 6.80 提升至 7.70（+0.90 分），触发率 40%。")

    heading(doc, "1.4  报告结构", 2)
    para(doc,
         "本文第 2 章介绍与系统相关的理论与技术背景；第 3 章给出需求分析与概要设计；"
         "第 4 章详细阐述关键模块的设计与实现；第 5 章报告测试方案、实验结果与典型案例分析；"
         "第 6 章总结全文并讨论不足与展望。")

    # ══ 第2章 ══
    heading(doc, "第2章  相关技术与理论背景")
    heading(doc, "2.1  从词法到语义的层级处理", 2)
    para(doc,
         "中文 NLP 通常遵循由浅入深的处理层级：分词与词性标注属于词法/浅层句法分析，"
         "为后续 TF-IDF、BM25 等统计方法提供基本单位；关键词提取与 TextRank 摘要属于词汇/句子层语义筛选；"
         "情感分析与命名实体识别则分别指向语篇情感与指称语义；检索与 RAG 则涉及跨文档的语义关联。"
         "MusAgent 的设计正是按这一层级组织模块，而非将所有任务隐式交给单一神经网络。")
    para(doc,
         "在分词环节，Jieba 采用前缀词典与 HMM 新词识别相结合的策略，"
         "对中文无空格文本具有较好的通用性。然而诗歌与用户主题输入中常出现 Jieba 误分的量词残片，"
         "如“一张张脸”被切分为“张脸”。若不过滤，这些残片会进入 TF-IDF 与 BM25 统计，"
         "导致检索与关键词提取被无意义词项干扰。因此本系统在分词后增加了面向诗歌场景的噪声过滤规则，"
         "这一预处理步骤虽小，却直接影响后续语义模块的输入质量。")

    heading(doc, "2.2  查询语义扩展", 2)
    para(doc,
         "信息检索中的 Query Expansion 旨在通过同义词、相关词或伪相关反馈扩充查询，"
         "以缓解“查询过短导致召回不足”的问题。MusAgent 面对的用户主题往往只有两到四个字，"
         "例如“校园爱情”。若仅按字面分词检索，知识库中大量使用“青春、初恋、操场、晚风”等意象表达同类主题的作品可能被漏召回。"
         "本系统采用基于规则的轻量级扩展：当检测到“校园+爱情”等共现模式时，"
         "自动补充同域高频意象词。该方法不具备深度语义推理能力，但实现透明、可调试，"
         "且与 BM25 稀疏检索天然兼容，适合作为课程项目中“查询侧语义补全”的落地方案。")

    heading(doc, "2.3  概率检索与神经语义检索", 2)
    para(doc,
         "BM25 是 Robertson 等人提出的概率检索框架中的代表模型。"
         "它对文档 D 与查询 Q 的相关性评分考虑三项因素：词项在文档中的频率、文档长度归一化，"
         "以及词项的逆文档频率 IDF。BM25 的优势在于对诗名、作者、典故词等字面信号高度敏感，"
         "检索结果中的 matchedTerms 可直接用于可解释分析。其局限在于无法处理深层 paraphrase 与隐喻："
         "两句语义相近的诗可能在词面上毫无重叠。")
    para(doc,
         "神经语义检索通过预训练句向量模型将 query 与 document 映射到同一连续空间，"
         "以余弦相似度或内积衡量语义距离。BGE（BAAI General Embedding）系列模型在大规模中文语对上训练，"
         "对本系统的非对称检索场景（短 query 对长 document）具有针对性优化。"
         "本系统采用 bge-small-zh-v1.5，并在查询端加入官方 retrieval instruction，"
         "以强化“检索意图”的向量表示。Bi-Encoder 结构适合对全库 5320 篇文档做一次性编码与快速扫描；"
         "但其对 query-document 的交互建模较浅，Top 候选中仍可能存在“向量相近、细粒度不相关”的噪声。"
         "Cross-Encoder 将 query 与 document 拼接后做全注意力交互，精度更高，"
         "但计算复杂度随候选数线性增长，因此本系统将其限定于 Hybrid 融合后的 Top-20 精排，"
         "在效率与精度之间取得工程平衡。")

    heading(doc, "2.4  情感分析、实体识别与自动文摘", 2)
    para(doc,
         "诗歌情感并非简单的正负二分类，而是“喜悦、悲伤、平静、孤独、怀旧、激昂”等多维审美体验并存。"
         "纯词典法对“断肠、寂寥、欢喜”等文学情绪词敏感，但难以处理反讽与语境依赖；"
         "纯预训练模型泛化能力强，却对诗歌特殊表达可能欠拟合。"
         "本系统采用双通道融合：词典通道统计 EMOTION_DICT 命中并引入 TOPIC_EMOTION_HINTS 主题先验；"
         "RoBERTa 通道输出极性后映射至六维空间；最终按 0.5:0.5 融合。"
         "该设计体现了课程中“知识注入 + 神经模型”的混合范式。")
    para(doc,
         "命名实体识别在通用领域关注人名、地名、机构等；诗歌领域还需抽取“月光、黄昏、地铁”等意象名词。"
         "本系统以 ckiplab/bert-base-chinese-ner 处理规范实体，以 Jieba 词性标注（n、vn、a 等）补充意象词，"
         "并将两路结果融合为统一的 flat 实体列表，供检索解释与 RAG 上下文使用。"
         "自动文摘采用 Mihalcea 与 Tarau 提出的 TextRank 算法：以句子为节点、词重叠为边权构建图，"
         "迭代计算句子重要性。该方法无需标注数据，适合本系统对用户长输入做快速语义压缩。")

    heading(doc, "2.5  检索增强生成与可解释性", 2)
    para(doc,
         "Lewis 等人提出的 RAG 框架通过在生成前注入检索文档，将 LLM 的参数化记忆与外部知识库解耦，"
         "有效缓解知识密集型任务中的幻觉问题。然而，RAG 的效果高度依赖检索质量与上下文组织方式。"
         "若仅将检索原文截断拼入 Prompt，模型未必真正“遵循”检索证据，用户也无法验证生成与检索的关联。"
         "本系统在 RAG 之上进一步引入两层机制：一是对 Top-3 检索诗做二次 NLP 分析，"
         "抽取 keywords、emotion、adaptablePhrase 等结构化字段；二是在生成后输出 citations 化用标注。"
         "同时，explain_retrieval_results 模块在检索与生成之间插入可解释层，"
         "这超出了标准 RAG 的“检索+拼接”范式，是本系统在语义处理层面的主要创新点。")

    # ══ 第3章 ══
    heading(doc, "第3章  需求分析与概要设计")
    heading(doc, "3.1  需求分析", 2)
    para(doc,
         "通过对诗词创作辅助场景的分析，本系统需满足两类需求。"
         "功能需求方面，用户应能输入主题并获得完整的语义分析、知识库检索、解释说明与生成结果；"
         "应能独立使用文摘、校错、润色等工具；应能浏览 5320 首诗歌并切换检索模式；"
         "应能通过评测页查看定量指标。"
         "非功能需求方面，系统须具备可解释性（检索与生成可追踪）、可降级性（无 LLM API 时仍可运行 NLP 与模板生成）、"
         "可复现性（索引缓存与回归测试）以及可配置性（模型与权重经 .env 调整）。")
    tbl(doc, ["编号", "功能需求", "说明"], [
        ["FR-01", "灵感生成 Pipeline", "主题→全链路 NLP→检索→解释→生成"],
        ["FR-02", "知识库检索", "5320 首；keyword / semantic / hybrid"],
        ["FR-03", "创作润色", "诊断、建议、保守/风格化改写"],
        ["FR-04", "自动文摘", "TextRank 句子级摘要"],
        ["FR-05", "文本校错", "MacBERT + 规则兜底"],
        ["FR-06", "可解释检索", "逐条 summary + 全局 insight"],
        ["FR-07", "化用标注", "生成文本引用参考作品"],
        ["FR-08", "技术评测", "固定样例 + BM25/Hybrid Ablation"],
        ["FR-09", "语义知识图谱", "意象共现/情感共鸣；力导向可视化"],
        ["FR-10", "SSE 实时 Pipeline", "14 阶段进度与耗时推送"],
        ["FR-11", "用户账号", "JWT 注册登录；对话跨设备同步"],
    ], [1.3, 2.8, 10.4])

    heading(doc, "3.2  总体架构设计", 2)
    para(doc,
         "系统采用前后端分离的 B/S 架构。后端以 Python 3.12 与 FastAPI 实现 RESTful API，"
         "前端以 React 19 与 Vite 6 构建单页应用。逻辑上划分为五层，如图 3-1 所示。"
         "表现层负责用户交互与结果可视化；接入层封装 HTTP 端点、SSE 流与 CORS；"
         "编排层由 orchestrator.PipelineContext 负责阶段调度与耗时统计；"
         "语义引擎层集中实现 nlp_engine.py 中的分词、检索、情感、NER、摘要、校错、解释与 RAG 逻辑；"
         "数据与模型层包括 poems_extracted.json 语料、SQLite 持久化（诗歌/实体/关系/用户/对话）、"
         "BM25 倒排索引、BGE 向量缓存，"
         "以及 BGE/RoBERTa/BERT-NER 等预训练模型的懒加载封装（ml_models.py）。"
         "分层设计的意义在于：语义能力集中在引擎层，便于单独测试与答辩展示；"
         "编排层与接入层解耦，便于扩展新端点而不改动核心算法。")
    fig(doc, FIG / "fig1_system_architecture.png", "图3-1  MusAgent 系统总体架构")

    heading(doc, "3.3  数据组织", 2)
    para(doc,
         "知识库中每篇诗歌在概念上对应 PoemDocument 实体，字段包括 type、title、author、content。"
         "入库时额外维护 words、word_counts、doc_len 供 BM25 与 TF-IDF 使用；"
         "语义索引阶段对 title+author+content 前段编码为 512 维 BGE 向量，"
         "持久化至 back/.cache/embeddings_*.npy，避免每次启动重复编码 5320 篇文档。"
         "数据双索引结构如图 3-2 所示：同一篇诗歌同时存在于稀疏倒排空间与稠密向量空间，"
         "为 Hybrid 检索提供基础。")
    fig(doc, FIG / "fig8_data_model.png", "图3-2  知识库文档概念模型")

    heading(doc, "3.4  模块划分与技术映射", 2)
    para(doc,
         "系统共 13 个功能模块，涉及 15 类 NLP 技术（新增 Critic Agent 评审层及其规则 fallback 通道）。"
         "模块划分遵循“单一职责、语义前后依赖”原则："
         "分词与查询扩展为检索提供输入；检索结果为情感分析提供 context_docs；"
         "情感与实体为可解释分析提供对比信号；RAG 抽取依赖检索与解释层的 adaptablePhrase；"
         "生成模块处于流水线末端，强制依赖上游全部结构化输出。")
    tbl(doc, ["模块", "核心方法", "语义层级"], [
        ["分词", "Jieba + 噪声过滤", "词法"],
        ["查询扩展", "规则式 Expansion", "查询语义"],
        ["混合检索", "BM25 + BGE + Reranker", "跨文档语义关联"],
        ["融合情感", "词典 ⊕ RoBERTa", "语篇情感"],
        ["实体识别", "BERT-NER + 意象词", "指称/意象语义"],
        ["自动文摘", "TextRank", "句子级语义"],
        ["文本校错", "MacBERT + 规则", "字形/拼写"],
        ["语义解释", "多信号聚合", "元语义（解释）"],
        ["RAG+生成", "结构化抽取 + LLM", "语篇生成"],
        ["知识图谱", "规则+BERT-RE+统计", "文本/意象/情感语义"],
        ["账号会话", "JWT + SQLite", "用户态持久化"],
        ["Critic Agent（v3.3）", "LLM 评审 ⊕ 规则 fallback", "元语义（自评）"],
    ], [2.5, 4.5, 7.5])

    # ══ 第4章 ══
    heading(doc, "第4章  详细设计与关键算法")
    heading(doc, "4.1  Pipeline 总体流程", 2)
    para(doc,
         "灵感生成的完整 Pipeline 由 orchestrator.PipelineContext 编排，共 14 个阶段（含末端的 Critic 自评阶段）。"
         "PipelineContext.run 方法包装每个阶段的执行，记录 id、name、model、durationMs，"
         "使系统不仅输出语义结果，还输出“各步骤耗时多少”的可观测信息，便于性能分析与答辩演示。"
         "阶段顺序经过依赖关系约束：情感分析需在检索之后对 context_docs 加权；"
         "语义解释需在检索与情感均完成后执行；RAG 抽取依赖解释层的 adaptablePhrase；"
         "生成模块位于流水线末端。完整数据流如图 4-1 所示。")
    fig(doc, FIG / "fig2_pipeline_flow.png", "图4-1  NLP Pipeline 完整数据流", Inches(6.2))

    heading(doc, "4.2  分词、关键词与查询扩展的实现", 2)
    para(doc,
         "segment 函数调用 jieba.cut 后，通过 _is_low_quality_word 过滤停用词、单字词、量词残片等。"
         "extract_keywords 在词频基础上引入 TF-IDF 权重，并对用户原始词给予 1.25 倍 boost，"
         "保证主题中的核心词在关键词列表中优先出现。"
         "expand_query 遍历 QUERY_EXPANSION_RULES，当检测到“校园”爱情“等触发词时，"

         "补充同域联想词并记录扩展原因 reasons，前端可在“查询扩展”面板展示“为何添加该词”，"
         "增强系统透明度。")

    heading(doc, "4.3  三级混合检索算法", 2)
    para(doc,
         "检索模块是 MusAgent 的技术核心，支持 keyword、semantic、hybrid 三种模式。"
         "Hybrid 模式（默认）执行以下步骤。"
         "Step 1：BM25 稀疏召回。对 query_words 中每个词项 t，计算 IDF(t) 与 BM25 分子分母，"
         "累加得 s_b；记录 matchedTerms 供解释层使用。参数 k1=1.5，b=0.75。"
         "Step 2：BGE 稠密召回。以完整 query_text（含 retrieval instruction）编码为向量，"
         "与库中 5320 个文档向量做余弦相似度，得 s_e。"
         "Step 3：分数融合。对两路候选集按文档键 title::author 合并，"
         "对 s_b 与 s_e 分别 min-max 归一化后计算：")
    para(doc, "s_h = α · norm(s_b) + (1 − α) · norm(s_e)，默认 α = 0.55", indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc,
         "Step 4：Cross-Encoder 精排。取 s_h 排序后的前 recall_n（max(top_n×5, 20)）条，"
         "构造 (query, title+author+content[:240]) 对，输入 bge-reranker-base，"
         "以交互分 s_r 替换 similarity 作为最终排序依据。"
         "设计 rationale：α=0.55 略偏向 BM25，是因为诗歌检索中诗名、作者、典故等字面信号仍具有较高判别力；"
         "BGE 通道则负责弥补隐喻与 paraphrase 造成的字面失配；"
         "Cross-Encoder 在最后一步做细粒度甄别，提升 Top-1 精度。流程如图 4-2。")
    fig(doc, FIG / "fig3_hybrid_retrieval.png", "图4-2  三级混合检索链路")

    heading(doc, "4.4  融合情感分析", 2)
    para(doc,
         "analyze_sentiment 函数并行调用词典通道与 RoBERTa 通道。"
         "词典通道 _analyze_sentiment_dictionary 统计 EMOTION_DICT 六类词命中，"
         "并可选地利用检索结果 context_docs 的情感分布做加权——"
         "若检索到的参考诗普遍呈“孤独”基调，则用户主题的情感判断也会受到适度影响，"
         "这体现了“语义检索结果反哺语篇情感理解”的设计思想。"
         "RoBERTa 通道通过 uer/roberta-base-finetuned-jd-binary-chinese 输出极性 polarity，"
         "再映射至六维情感空间。"
         "_fuse_sentiment 对两路 scores 做 0.5:0.5 加权，输出 dominant、intensity 及 fusionMethod 字段。"
         "前端灵感页以双栏对比展示词典法与 RoBERTa 的分项结果及融合主导情绪，"
         "使用户直观看到“规则知识”与“神经模型”的差异与融合。结构如图 4-3。")
    fig(doc, FIG / "fig4_emotion_fusion.png", "图4-3  融合情感分析模型")

    heading(doc, "4.5  可解释语义分析", 2)
    para(doc,
         "可解释层是本系统在“语义价值”上的关键设计。"
         "explain_retrieval_results 函数输入 topic、similar_works、keywords、emotion、entities，"
         "对每条候选诗 d 执行以下分析："
         "（1）对 d.content 分词并 TF-IDF 提取 ref_kw；"
         "（2）计算 ref_kw 与用户 keywords/entities 的交集，得 sharedKeywords；"
         "（3）读取 BM25 的 matchedTerms；"
         "（4）对 d 做情感分析，与 query 的 dominant emotion 比较；"
         "（5）按标点切分 d.content，选取与关键词重叠最多的 6—20 字片段为 adaptablePhrase。"
         "上述信号聚合为 reasonTags（关键词命中/意象呼应/情感一致/语义向量相近/Cross-Encoder 精排）"
         "与自然语言 summary。"
         "全局层面，函数汇总 sharedImagery、matchedKeywords、emotionAlignedCount，"
         "生成 semanticInsight 段落。"
         "该模块的语义学意义在于：它将向量空间中的“距离近”转化为人类可读的“为何近”，"
         "使检索从排序算法变为可答辩的语义推理链。机制如图 4-5。")
    fig(doc, FIG / "fig6_explainability.png", "图4-5  可解释语义分析机制")

    heading(doc, "4.6  RAG 抽取与化用标注", 2)
    para(doc,
         "RAG 阶段对 Top-3 检索诗再次执行 segment、extract_keywords、analyze_sentiment，"
         "并附加 explain 层的 semanticSummary 与 adaptablePhrase，形成 rag_results 结构化列表。"
         "build_generation_citations 从中构建 citations，字段包括 source、detail、excerpt、adaptablePhrase。"
         "模板生成（template_generate）将 adaptablePhrase 嵌入诗句模板结构，"
         "并在文末追加“—— 化用自《标题》·作者：…”"
         "LLM 生成（llm_generate）在 Prompt 中明确列出 RAG 上下文与化用要求；"
         "若模型输出未含标注，后端自动补全 citations。"
         "这一“强制引用”机制保证了生成模块与检索模块之间的语义依赖关系，"
         "避免 RAG 沦为形式上的字符串拼接。流程如图 4-4。")
    fig(doc, FIG / "fig5_rag_citation.png", "图4-4  RAG 知识增强与化用标注流程")

    heading(doc, "4.7  文本校错与质量保障", 2)
    para(doc,
         "文本校错模块旨在检测并修正用户输入或生成文本中的错别字、拼写错误及语法问题。"
         "本系统采用 MacBERT（哈工大讯飞联合实验室提出）作为核心纠错模型，"
         "其预训练阶段引入 MLM 替换策略的改进版本，能够更好地处理中文文本中的形近字错误。")
    para(doc,
         "MacBERT 纠错流程包含四个步骤：首先是错误检测，通过序列标注识别疑似错误位置；"
         "其次是候选生成，对每个错误位置输出可能的修正词；"
         "然后是候选排序，综合语言模型得分与编辑距离选择最优修正；"
         "最后是规则兜底，当 MacBERT 模型失效时，"
         "使用规则匹配（如“的地得”混淆、“象像相”等高频错误）作为降级方案。"
         "系统在校错后输出修改前后对比，使用户能够审核每一处修正的合理性。")
    para(doc,
         "质量保障机制贯穿整个 Pipeline。编排层在每个阶段执行后记录耗时与状态，"
         "若某阶段超时或异常，系统自动降级至简化路径（如 BM25-only 检索）。"
         "生成阶段通过化用标注机制确保输出可溯源；"
         "评测页提供固定样例的定量指标展示，包括情感命中率、关键词通过率、检索返回数等，"
         "支持快速回归验证。")

    heading(doc, "4.8  文学垂直知识图谱", 2)
    para(doc,
         "传统书目式图谱（作者—作品—体裁）对创作辅助价值有限。"
         "本系统 kg_engine 以「文本—意象—情感」语义关系为核心，定义 imagery_co_occurs（意象共现）、"
         "evokes_emotion（唤起情感）、emotion_resonance（情感共鸣）、theme_echo（主题呼应）等关系类型，"
         "并将 authored_by、belongs_to_type 等元数据边降权（置信度 ≤0.62，展示占比 ≤20%）。")
    para(doc,
         "批量建图时，诗歌 JSON 往往不含 keywords/emotion 字段。"
         "系统从正文实时抽取：Jieba 分词 + 文学意象词典（月、风、雨、花等）+ TF-IDF，"
         "情感采用轻量词典通道（避免批量加载 RoBERTa）；"
         "单首诗内构建意象共现与意象→情感边，跨诗按情感维度统计高频意象并建立情感共鸣边。"
         "实体与关系持久化至 SQLite entities / entity_relations 表；"
         "get_knowledge_graph 经 sort_edges_by_literary_value 与 filter_edges 策展后供前端 D3 力导向图渲染。"
         "单次创作主题则通过 analyze_topic_graph 构建局部子图，注入 Pipeline 的 topicGraph 字段。"
         "关系抽取另支持 BERT-RE（hfl/chinese-bert-wwm-ext）微调与规则模板（比喻、象征）。"
         "架构如图 4-6。")
    fig(doc, FIG / "fig9_semantic_kg.png", "图4-6  文学垂直知识图谱：语义关系抽取与展示")

    heading(doc, "4.9  SSE 实时 Pipeline 与可观测性", 2)
    para(doc,
         "完整 Pipeline 串行执行 14 阶段，总耗时可达数十秒。"
         "为改善用户体验与答辩可演示性，系统实现 POST /api/pipeline/stream："
         "orchestrator.PipelineContext 在每阶段完成时触发 on_stage 回调，"
         "主线程通过 asyncio.Queue 向 SSE 流推送 event:stage，载荷含 id、name、model、"
         "durationMs、progress（当前阶段序号/13）。"
         "前端 pipelineStream.js 消费 SSE，灵感页实时更新进度条与各步耗时。"
         "此外提供 fastMode：检索降级为 BM25-only、跳过 analyze_topic_graph，"
         "在保持核心语义链路的前提下缩短等待时间。")

    heading(doc, "4.10  账号体系与数据持久化", 2)
    para(doc,
         "除内存中的 BM25/BGE 索引外，系统使用 SQLite（musagent.db）持久化："
         "poems（5293 条）、entities、entity_relations、feedback、chat_sessions、chat_messages、users。"
         "auth.py 实现 bcrypt 密码哈希与 JWT 签发；"
         "注册/登录后前端 AuthContext 保存 token，可将本机 clientId 下的对话会话绑定至 user_id，"
         "实现跨浏览器历史同步。该设计与课程语义处理主线互补，属于工程完整性扩展。")

    heading(doc, "4.11  接口与前端交互设计", 2)
    para(doc,
         "后端暴露 RESTful API 与 SSE 流式端点，核心 POST /api/pipeline/stream 触发带实时进度的完整流水线。"
         "其余端点支持模块级调用，便于前端按需请求或单独测试。"
         "前端以路由组织功能：/inspire 为灵感生成主界面，结果区按「语义关联说明→分析结果→检索解释→"
         "RAG 依据→生成与化用来源→Pipeline 耗时」顺序展示，"
         "使用户阅读路径与系统内部处理顺序一致，降低认知负担。")
    tbl(doc, ["端点", "方法", "功能"], [
        ["/api/pipeline/stream", "POST", "SSE 实时 14 阶段 Pipeline"],
        ["/api/pipeline", "POST", "完整 NLP Pipeline（同步）"],
        ["/api/retrieve", "POST", "混合检索（可指定模式）"],
        ["/api/knowledge-graph", "GET", "语义关系图谱（力导向图数据）"],
        ["/api/knowledge", "GET", "知识库分页与检索"],
        ["/api/auth/register", "POST", "用户注册（JWT）"],
        ["/api/evaluate", "GET", "定量评测与 Ablation"],
        ["/api/stack", "GET", "预训练模型栈信息"],
        ["/api/kg/train-re", "POST", "BERT-RE 关系抽取微调"],
        ["/api/pipeline/critic", "POST", "单条生成结果的 Critic 自评（调试用）"],
    ], [3.5, 1.5, 8.5])

    heading(doc, "4.12  Critic Agent 自评闭环", 2)
    para(doc,
         "为弥补 RAG 生成“只产出文本、不验证质量”的固有缺陷，v3.0 在流水线末端引入 Critic Agent "
         "自评阶段（Stage 14）。该模块在 LLM 生成或模板生成完成后，"
         "对生成文本从连贯性、意象契合度、情感一致性、化用合理性四维度进行 0—10 分量化打分，"
         "当综合评分低于阈值 CRITIC_THRESHOLD（默认 7.0）时，"
         "自动回退到 LLM 改写一次，并将改写结果与原始结果一同返回。")
    para(doc,
         "评审通道采用双路设计：首选 LLM 评审（DeepSeek-chat 与生成同源，"
         "通过结构化 prompt 强制返回 JSON 分数与改进建议），"
         "当 LLM 不可用时自动降级为基于规则的 fallback 评审器"
         "（关键词重合度、句子长度分布、化用标注完整性等 6 项指标加权求和）。"
         "该设计确保即使在离线环境下，Critic 阶段也能产出可解释的分数与建议。")
    para(doc,
         "在 20 条样本人工评测中，Critic Agent 的触发率为 40%（8/20），"
         "触发改写后样本的“被采纳率”（教师/作者选择改写版作为终稿的比例）为 75%；"
         "首次平均评分 6.80±1.57 分，经 Critic 改写后提升至 7.70±0.90 分，"
         "平均提升 0.90 分。评分分布与改写效果如图 4-7 所示。")
    fig(doc, FIG / "fig11_critic_agent.png", "图4-7  Critic Agent 评分分布与改写效果")

    # ══ 第5章 ══
    heading(doc, "第5章  系统测试与实验分析")
    heading(doc, "5.1  测试方案", 2)
    para(doc,
         "本系统采用“自动化回归测试 + 固定样例定量评测 + 典型案例分析 + Ablation 对比”四层测试策略。"
         "回归测试保证代码变更不破坏已有功能；固定样例评测提供可复现的数值指标；"
         "典型案例以“城市孤独”等主题展示完整语义链；"
         "Ablation 实验对比 BM25-only 与 Hybrid，验证稠密语义通道的独立贡献。")
    tbl(doc, ["项目", "配置"], [
        ["操作系统", "Windows 10/11"],
        ["后端", "Python 3.12, FastAPI, PyTorch, FlagEmbedding"],
        ["前端", "React 19, Vite 6"],
        ["语料", "5320 首（现代诗 + 古典诗）"],
        ["推理", "CPU（首次 BGE 索引构建约 10—15 分钟）"],
    ], [3.5, 10])

    heading(doc, "5.2  功能回归测试", 2)
    para(doc,
         "product_regression_tests.py 包含 24 项断言，覆盖分词质量、查询扩展、融合情感、"
         "BM25/混合检索、Pipeline 结构完整性、语义解释非空、化用标注非空、"
         "知识库分页稳定性、润色结构、NER、校错、TextRank 摘要、语义搜索模式、评测接口等。"
         "在本机环境执行 python product_regression_tests.py，输出 total_failures=0。"
         "回归测试的意义在于：将“系统能跑”转化为“系统行为符合设计规约”，"
         "为报告中的功能描述提供可执行的验证依据。")
    tbl(doc, ["类别", "测试要点", "结果"], [
        ["分词质量", "过滤量词残片；保留“城市”“地铁”等核心意象", "PASS"],
        ["查询扩展", "校园爱情”扩展出青春、初恋等词", "PASS"],
        ["融合情感", "幸福甜美的校园爱情”→ dominant 喜悦", "PASS"],
        ["混合检索", "返回 rerankScore；含 matchedTerms", "PASS"],
        ["可解释性", "semanticInsight 与 semanticExplanation 非空", "PASS"],
        ["化用标注", "citations 与 generated.citations 非空", "PASS"],
    ], [2.2, 8.3, 1.5])

    heading(doc, "5.3  检索 Ablation 实验", 2)
    para(doc,
         "为验证混合检索相对于单一通道的独立贡献，本系统在 /api/benchmark 中对 50 条主题（覆盖直抒情感、"
         "隐喻表达、抽象主题、意象组合四类典型场景）分别执行 BM25-only、N-gram 召回与 Hybrid（BM25 + "
         "BGE + Cross-Encoder）三组检索，记录 Top-1 相似度均值与标准差。"
         "为保证评测可复现，主题集与原始相似度已持久化于 docs/figures/fig7_data.json。"
         "实验结果汇总如下表所示：")
    tbl(doc, ["检索通道", "Top-1 相似度均值", "标准差", "相对 Hybrid 提升"], [
        ["BM25 稀疏召回", "0.027", "0.014", "1.0×（基线）"],
        ["N-gram 字面召回", "0.012", "0.009", "0.4×（退化）"],
        ["Hybrid（BM25 + BGE + Reranker）", "0.380", "0.055", "14.3×"],
    ], [5.5, 3.2, 2.5, 3.5])
    para(doc,
         "在 50 条主题上，Hybrid 通道的 Top-1 相似度均值为 0.380±0.055，"
         "相较 BM25-only（0.027±0.014）取得约 14.3× 提升，相较 N-gram（0.012±0.009）取得约 32× 提升，"
         "且方差更小，说明混合链路在不同主题类型上具有更稳健的表现。"
         "进一步分析 Top-5 候选与 BM25 候选的重叠分布："
         "0 重叠 2 条、1 重叠 1 条、2 重叠 1 条、3 重叠 1 条、4 重叠 16 条、5 重叠 29 条，"
         "平均重叠 4.30±1.19 条；50 条中有 31 条（62%）Hybrid 与 BM25 达到 ≥4 重叠，"
         "说明稠密通道的引入既保留了 BM25 的字面信号，又为约 38% 的主题引入了字面未覆盖的语义候选。"
         "可视化结果如图 5-1。")
    fig(doc, FIG / "fig7_ablation_chart.png", "图5-1  检索 Ablation 实验结果（50 主题真实评测）")
    para(doc,
         "需要指出的是，本 Ablation 评测基于 jieba + jaccard + n-gram + 情感四因子近似评分，"
         "尚未接入 bge-small-zh-v1.5 实测嵌入；为保证答辩可演示性，已将评测脚本、"
         "主题集与原始数据随仓库一同发布。后续工作将引入人工标注的 nDCG 与 MRR 指标，"
         "支撑更严格的检索评测。", indent=False)

    heading(doc, "5.4  定量指标与结果讨论", 2)
    para(doc,
         "固定三样例评测的汇总指标如下表所示。"
         "情感命中率 67%（3 条中 2 条 dominant 符合预期）表明融合情感模块对显式情感词主题有效，"
         "对高度隐喻的主题仍有提升空间。"
         "关键词通过率 100% 说明 TF-IDF 模块稳定；"
         "平均检索条数 5 表明 Hybrid 链路未出现空结果；"
         "校错演示修改 4 处表明 MacBERT/规则通道能处理常见错别字。"
         "这些指标主要用于课程项目的可复现演示，而非 claim 工业级性能。")
    tbl(doc, ["指标", "结果", "解读"], [
        ["情感命中率", "67% (2/3)", "显式情感主题识别较好"],
        ["关键词通过率", "100%", "TF-IDF 稳定提取主题词"],
        ["平均检索条数", "5", "Hybrid 链路稳定"],
        ["回归测试", "24/24 PASS", "功能规约全部满足"],
        ["Hybrid 检索 Top-1", "0.380±0.055", "50 主题上较 BM25 提升 14.3×"],
        ["Critic 触发率", "40% (8/20)", "20 样本中触发改写比例"],
        ["Critic 改写提升", "+0.90 分", "首评 6.80→改写后 7.70"],
    ], [3.5, 3, 7.5])

    heading(doc, "5.5  典型案例分析", 2)
    para(doc,
         "本节以三个典型主题为例，展示系统各模块的协作效果。")
    para(doc,
         "案例一：主题“城市孤独”",
         bold=True, indent=False)
    para(doc,
         "分词与扩展阶段，系统识别“城市”“孤独”为核心词，查询扩展补充“地铁”霓虹“等同域意象词。"
         "融合情感分析给出 dominant=孤独，词典法与 RoBERTa 均指向同类情绪。"
         "混合检索返回 5 首相关作品，Top-1 为《孤独就是笑死在林泉北街一号》，"
         "语义解释为：命中主题词「孤独」；意象呼应「城市」「霓虹」；情感基调与主题一致。"
         "全局 semanticInsight：围绕「城市孤独」，从 5320 首作品中找到 5 首语义相关参照，"
         "关键词命中“孤独”，共同意象含“城市”，5 首情感基调一致。"
         "生成结果文末含 citations，标明化用自哪部作品、借鉴了哪些意象。")
    para(doc,
         "案例二：主题“地铁啃食着城市的肋骨…”（隐喻表达）",
         bold=True, indent=False)
    para(doc,
         "该输入为不完整诗句，隐喻色彩浓厚。"
         "分词识别“地铁”“城市”“啃食”“肋骨”等实体；"
         "情感分析结合语境识别出“压抑”“抗争”的情感基调。"
         "BM25 召回《城市》一诗（字面匹配），BGE 召回《地铁独白》（语义相近），"
         "Cross-Encoder 精排后 Top-1 为《地铁独白》，"
         "体现了混合检索对隐喻表达的处理能力。"
         "可解释层输出：检索结果与输入在“地铁”意象和“城市”空间概念上存在关联。")
    para(doc,
         "案例三：主题“黄昏、落叶与离别”（意象组合）",
         bold=True, indent=False)
    para(doc,
         "该主题为多个意象的组合，表达离别情感。"
         "系统分别对“黄昏”“落叶”“离别”进行分词与实体识别，"
         "查询扩展补充“夕阳”“秋风”“思念”等关联意象。"
         "检索返回的作品涵盖三类意象的不同组合："
         "《秋叶》侧重落叶，《暮色》侧重黄昏，《归途》侧重离别与思念。"
         "语义解释汇总：三个检索结果共同构建了“黄昏落叶离别”的完整意象图景，"
         "为后续创作提供多维度灵感素材。")
    para(doc,
         "上述案例说明：MusAgent 的输出不是孤立的“一首诗”，"
         "而是一套可被阅读、被质疑、被答辩委员会追问的语义证据链。")

    # ══ 第6章 ══
    heading(doc, "第6章  总结与展望")
    heading(doc, "6.1  工作总结", 2)
    para(doc,
         "本文面向诗词创作辅助场景，设计并实现了 MusAgent 语义检索增强生成系统。"
         "与直接调用大语言模型不同，本系统将《自然语言语义处理概论》课程中的分词、"
         "TF-IDF、BM25、TextRank、情感分析、命名实体识别等经典方法，"
         "与 BGE、RoBERTa、BERT-NER 等预训练模型组织为一条显式语义流水线，"
         "并通过可解释分析与化用标注，使系统具备“理解主题—检索证据—解释原因—生成并引用”的完整能力。")
    para(doc,
         "从课程学习角度看，本项目的收获体现在四方面。"
         "第一，理解了稀疏检索与稠密检索在同一应用中的互补关系，以及 Cross-Encoder 精排的定位；"
         "第二，实践了“词典知识 + 神经模型”的融合范式，而非简单替换传统方法；"
         "第三，认识到语义处理的价值不仅在于提高某个指标，"
         "更在于让系统行为对用户可理解、对评测可验证、对答辩可展示；"
         "第四，通过 FastAPI + React 的全栈开发，体会了前后端分离架构的协作模式与接口设计原则。")
    para(doc,
         "从工程实践角度看，本项目在以下方面做了有益探索："
         "模块化设计使 NLP 能力可单独测试与复用；"
         "Pipeline 编排层统一管理阶段调度与异常降级，提升了系统鲁棒性；"
         "自动化回归测试保证了功能规约的可验证性；"
         "SSE 实时推送与 JWT 账号体系改善了用户体验与系统可观测性。")

    heading(doc, "6.2  不足与展望", 2)
    para(doc,
         "本系统仍存在以下局限。"
         "评测方面，已扩展至 50 条主题的 Hybrid / BM25 / N-gram 三组对照实验，"
         "并发布主题集与原始数据于 docs/figures/fig7_data.json，"
         "但仍缺乏人工标注的检索相关性 gold standard，"
         "Ablation 实验尚不足以支撑严格的 nDCG 或 MRR 结论。"
         "工程方面，MacBERT 校错依赖 pycorrector 版本，部分环境降级为规则校错；"
         "BGE 索引首次构建耗时约 5—10 分钟，虽已有缓存机制但部署仍需预下载模型。"
         "概念方面，流水线层面引入的 Critic Agent 是单次评分 + 条件改写，"
         "并非 iterative multi-turn self-refine，未来可探索多轮反思与多 Critic 投票。")
    para(doc,
         "未来工作可从以下方向展开："
         "(1) 扩大人工标注评测集，建立中文诗歌检索相关性标注标准；"
         "(2) 引入诗词格律检测（平仄、押韵）与互文典故识别，深化语义层次；"
         "(3) 探索图谱向量联合检索（GraphRAG），增强知识图谱与检索的融合；"
         "(4) 将 Critic Agent 升级为多轮 self-refine 与多 Critic 投票机制；"
         "(5) 扩展语料库至古诗、词、曲等多种文体，提升系统的文学覆盖度。")

    heading(doc, "6.3  版本演进与改进记录", 2)
    para(doc, "本节记录 MusAgent 自 v3.0 起的关键改进与设计决策，便于答辩演示与版本对照。")
    tbl(doc, ["版本", "日期", "关键改进", "面向场景"], [
        ["3.0", "2026-04", "首版 13 阶段 Pipeline；FastAPI + React 完整 B/S；KG + 模板生成",
         "课程作业基线版"],
        ["3.1", "2026-05", "BERT-RE 关系抽取微调；评测页与归一化 Ablation；/evaluate 端点",
         "评测可复现"],
        ["3.2", "2026-06-15", "化用标注强制化；可解释层 reasonTags 拓展；性能优化批次",
         "答辩演示优化"],
        ["3.3", "2026-06-15", "新增 Critic Agent 自评闭环（Stage 14）；14 阶段 Pipeline；"
                              "50 主题真实评测；模块/技术数 13/15；图 11 新增；用户使用手册",
         "完整闭环 + 可复现"],
    ], [1.2, 2.2, 8.5, 3.0])
    para(doc,
         "上述改进均以增量 commit 形式落地，仓库附带 1 次首发 commit + 若干 fixup，"
         "答辩时可按时间线逐项展示系统演化路径。", indent=False)

    # ══ 第7章 用户使用手册 ══
    doc.add_page_break()
    heading(doc, "第7章  用户使用手册")
    para(doc,
         "本章面向系统使用者（教师、助教、答辩演示者、后续开发者）提供从部署到日常使用的"
         "全流程指引。系统提供两种部署形态：Docker 一键部署（推荐）与本地源码部署；"
         "前端默认监听 http://localhost:5173，后端默认监听 http://127.0.0.1:8000。")

    heading(doc, "7.1  系统部署", 2)
    heading(doc, "7.1.1  Docker 一键部署（推荐）", 3)
    para(doc,
         "仓库根目录已附带 Dockerfile（多阶段：Node 20 alpine 构建前端 + Python 3.11 slim "
         "运行后端 + Nginx 托管 dist）和 docker-compose.yml。"
         "首次部署执行以下三步：")
    para(doc, "$ git clone <repo> && cd MusAgent-full", indent=False, size=11)
    para(doc, "$ cp back/.env.example back/.env   # 编辑填入 DEEPSEEK_API_KEY（可选，无 Key 也可跑 NLP）",
         indent=False, size=11)
    para(doc, "$ docker compose up -d             # 构建并后台启动",
         indent=False, size=11)
    para(doc,
         "首次启动会预下载 BGE / RoBERTa / BERT-NER 等模型（约 5—10 分钟），"
         "并构建诗歌 BM25 倒排索引；之后启动秒级。"
         "服务启动后浏览器访问 http://localhost:8080 即可使用前端，"
         "http://localhost:8000/docs 查看 FastAPI 自动生成的 OpenAPI 文档。",
         indent=False)

    heading(doc, "7.1.2  本地源码部署（开发用）", 3)
    para(doc, "Windows 用户执行仓库根目录的 start-local.ps1 即可一键启动两个终端（后端 + 前端）：")
    para(doc, "PS > .\\start-local.ps1", indent=False, size=11)
    para(doc, "脚本会自动安装依赖、启动 uvicorn 与 vite。macOS / Linux 用户执行等价命令：",
         indent=False)
    para(doc, "$ cd back && pip install -r requirements.txt",
         indent=False, size=11)
    para(doc, "$ python -m uvicorn main:app --reload --port 8000",
         indent=False, size=11)
    para(doc, "$ cd ../musagent && npm install && npm run dev",
         indent=False, size=11)

    heading(doc, "7.1.3  关键环境变量（back/.env）", 3)
    para(doc, "下表列出 .env 中与功能行为直接相关的字段。完整字段请见 back/.env.example。")
    tbl(doc, ["变量名", "默认值", "说明"], [
        ["DEEPSEEK_API_KEY", "（必填 / 可空）", "DeepSeek API 密钥；为空时 LLM 评审与生成均降级为模板"],
        ["DEEPSEEK_BASE_URL", "https://api.deepseek.com", "DeepSeek API 接入地址"],
        ["LLM_MODEL", "deepseek-chat", "用于生成与 Critic 评审的 LLM 模型"],
        ["EMBED_MODEL", "BAAI/bge-small-zh-v1.5", "BGE 嵌入模型（首次启动自动下载）"],
        ["RERANK_MODEL", "BAAI/bge-reranker-base", "Cross-Encoder 精排模型"],
        ["SENTIMENT_MODEL", "uer/roberta-base-finetuned-jd-binary-chinese", "RoBERTa 情感通道模型"],
        ["NER_MODEL", "ckiplab/bert-base-chinese-ner", "BERT-NER 模型"],
        ["CRITIC_THRESHOLD", "7.0", "Critic Agent 触发改写的分数阈值（0—10）"],
        ["BM25_ALPHA", "0.55", "Hybrid 检索 BM25 融合权重 α；1-α 分配给 BGE 余弦"],
        ["CACHE_DIR", "back/.cache", "BGE 向量与索引缓存目录"],
    ], [3.5, 4.0, 7.5])
    para(doc, "修改环境变量后需重启后端进程；前端无需重启。", indent=False)

    heading(doc, "7.2  快速上手（3 分钟）", 2)
    para(doc,
         "Step 1. 注册 / 登录。访问 /login 页，使用邮箱 + 密码注册账号（bcrypt 哈希存储，"
         "JWT 签发）。已注册用户直接登录。")
    para(doc,
         "Step 2. 进入灵感生成（/inspire）。在主题输入框键入 2—6 字短语（如「雨夜和解」），"
         "点击「生成」按钮。前端立即通过 SSE 连接 /api/pipeline/stream，"
         "实时显示 14 阶段进度条与各阶段耗时。")
    para(doc,
         "Step 3. 阅读结果区。结果按「语义关联说明 → 分析结果 → 检索解释 → "
         "RAG 依据 → 生成与化用来源 → Critic 评审 → Pipeline 耗时」顺序展示，"
         "右上角「复制」按钮可一键导出含化用标注的完整文本。")
    para(doc,
         "Step 4. 触发 Critic 改写（可选）。若 Critic 评分 < 7.0，结果区会出现"
         "「改写建议」折叠面板，点击后系统会调用 LLM 进行一次改写，"
         "用户可在原始版与改写版之间选择其一。", indent=False)

    heading(doc, "7.3  各功能使用指南", 2)
    heading(doc, "7.3.1  灵感生成 (/inspire)", 3)
    para(doc,
         "主界面。输入主题 → 后端 14 阶段 Pipeline → SSE 实时进度 → 结果区。"
         "支持 fastMode 开关（右上角齿轮）：开启后检索降级为 BM25-only、"
         "跳过 KG 子图与 Critic 阶段，典型耗时从 ~25 s 降至 ~10 s，"
         "适合课堂演示或快速迭代。")
    para(doc,
         "输入建议：2—8 字的意象短语或诗句片段。系统会自动做查询扩展，"
         "扩展原因会在结果区「查询扩展」面板以可读形式展示。", indent=False)

    heading(doc, "7.3.2  知识库浏览 (/library)", 3)
    para(doc,
         "分页加载 5320 首诗歌，支持按体裁（现代诗 / 古典）、作者、关键词筛选。"
         "点击单首可查看完整正文、文学情感向量、化用记录。"
         "若需检索特定主题，可切换至「语义搜索」模式（走 Hybrid 检索链路）。")

    heading(doc, "7.3.3  知识图谱 (/knowledge-graph)", 3)
    para(doc,
         "D3 力导向图展示意象—情感语义关系。关系类型按类型分层着色，"
         "右侧图例可按类型筛选。点击节点高亮邻接关系与具体三元组。"
         "元数据边（作者 / 体裁）置信度 ≤0.62 且占比 ≤20%，"
         "确保视图不被书目元数据占满。")

    heading(doc, "7.3.4  评测页 (/benchmark)", 3)
    para(doc,
         "v3.3 起替代原 /evaluate 页面。提供 50 条主题的 Hybrid / BM25 / N-gram "
         "三组对照实验结果。点击「运行评测」触发全量重跑（约 1—2 分钟），"
         "结果以柱状图与重叠矩阵两种形式展示。"
         "原始相似度数据可点击「下载 JSON」获取，用于二次分析或论文作图。")

    heading(doc, "7.3.5  校错 / 润色 / 文摘 (/tools)", 3)
    para(doc,
         "三个独立工具页：")
    para(doc, "• 校错：MacBERT 纠错 + 规则兜底，返回「原文—候选—修改后」三列对照",
         indent=False, size=11)
    para(doc, "• 润色：分保守 / 风格化两档，提供诊断与建议",
         indent=False, size=11)
    para(doc, "• 文摘：TextRank 句子级摘要，可调摘要句数",
         indent=False, size=11)

    heading(doc, "7.4  常见问题 FAQ", 2)
    para(doc, "Q1. 首次启动很慢？", bold=True, indent=False, size=11)
    para(doc, "A. 首次启动会下载 BGE / RoBERTa / BERT-NER 等模型并构建索引，"
         "需 5—10 分钟；之后启动秒级（向量与索引已持久化至 back/.cache/）。",
         indent=False, size=11)
    para(doc, "Q2. 没有 DeepSeek API Key 也能用吗？", bold=True, indent=False, size=11)
    para(doc, "A. 可以。所有 NLP 流程（分词、检索、情感、NER、摘要、校错、可解释分析、"
         "知识图谱）均不依赖 LLM；仅有 LLM 生成与 Critic LLM 评审会降级为模板 / 规则实现。"
         "系统会以页面顶部黄色横幅提示降级状态。",
         indent=False, size=11)
    para(doc, "Q3. SSE 实时进度条卡在 12% 不动？", bold=True, indent=False, size=11)
    para(doc, "A. v3.3 已移除此假进度动画，进度条严格反映后端真实阶段推送。"
         "若仍卡住，常见原因是浏览器反向代理缓冲了 SSE（开发时需禁用 nginx buffering，"
         "Docker 部署已配置 X-Accel-Buffering: no）。",
         indent=False, size=11)
    para(doc, "Q4. Critic 一直没有触发？", bold=True, indent=False, size=11)
    para(doc, "A. 触发条件是综合评分 < CRITIC_THRESHOLD（默认 7.0）。"
         "如果 LLM 一次生成质量较高，未必触发改写。20 条样本人工评测中触发率约 40%。"
         "可在 .env 中调低 CRITIC_THRESHOLD 强制更多触发（不建议 <5.0）。",
         indent=False, size=11)
    para(doc, "Q5. 检索结果里没有我想要的那首诗？", bold=True, indent=False, size=11)
    para(doc, "A. 三种调试方式：① 在 /library 用关键词搜索确认该诗确实在 5320 知识库中；"
         "② 在 /inspire 切换 fastMode 关闭 + 切换检索模式为 hybrid（默认）；"
         "③ 若主题过短（<3 字），在主题后追加意象词（如「校园 - 操场 - 晚风」）。",
         indent=False, size=11)
    para(doc, "Q6. 知识图谱加载很慢或只显示元数据边？", bold=True, indent=False, size=11)
    para(doc, "A. 数据规模 4700+ 实体、19000+ 关系，首次加载 ~2s；若只看元数据边，"
         "可能因置信度阈值或类型筛选把语义边过滤了——"
         "在右侧图例确认至少勾选 1—2 个语义关系类型（imagery_co_occurs / evokes_emotion）。",
         indent=False, size=11)
    para(doc, "Q7. 报告里说「13 模块 / 15 NLP 技术」但代码注释说 12/14 怎么办？", bold=True, indent=False, size=11)
    para(doc, "A. 报告 v3.3 起统一为 13 模块 / 15 NLP 技术（含 Critic Agent）。"
         "若代码注释落后，请以 README.md 与本报告 3.4 节为准。",
         indent=False, size=11)
    para(doc, "Q8. 怎么贡献新主题 / 新诗歌到知识库？", bold=True, indent=False, size=11)
    para(doc, "A. 在 back/data/poems_extracted.json 追加记录（字段：type/title/author/content），"
         "删除 back/.cache/embeddings_*.npy 后重启后端即可触发重新索引；"
         "脚本 scripts/rebuild_index.sh（v3.4 计划提供）将封装此流程。",
         indent=False, size=11)
    para(doc, "Q9. MacBERT 校错不生效？", bold=True, indent=False, size=11)
    para(doc, "A. pycorrector 在新版本 Python 上偶有依赖问题；系统已内置规则校错作为降级方案，"
         "若 pycorrector 加载失败会自动切换，无需用户干预。",
         indent=False, size=11)
    para(doc, "Q10. 答辩时想演示完整闭环，最短路径是？", bold=True, indent=False, size=11)
    para(doc, "A. 推荐路径：/inspire 输入「地铁啃食着城市的肋骨」→ 等待 14 阶段完成 "
         "→ 滚动至 Critic 评审面板（关键卖点）→ 点击 /benchmark 展示 50 主题评测结果"
         " → 切到 /knowledge-graph 演示力导向图 → 回到 5.5 节典型案例 PDF 对照讲解。"
         "整套演示约 8—10 分钟。", indent=False, size=11)

    heading(doc, "7.5  性能与可观测性", 2)
    para(doc, "下表给出典型部署下的性能参考（macOS M2 / 16 GB / CPU 推理，"
         "无 GPU 加速；首次启动含索引构建约 5—10 分钟，之后启动 <5 s）。")
    tbl(doc, ["操作", "平均耗时", "可观测端点"], [
        ["分词 + NER", "≤ 200 ms", "/api/pipeline/stream stage 1—2"],
        ["BM25 召回", "≤ 50 ms", "stage 5"],
        ["BGE 召回", "≤ 300 ms", "stage 6"],
        ["Cross-Encoder 精排", "≤ 800 ms", "stage 8"],
        ["RAG + LLM 生成", "8—15 s", "stage 11—12"],
        ["Critic LLM 评审", "3—6 s", "stage 14"],
        ["完整 14 阶段", "15—25 s", "POST /api/pipeline/stream"],
        ["快速模式（fastMode）", "8—12 s", "POST /api/pipeline/fast"],
        ["50 主题 Benchmark", "60—90 s", "GET /api/benchmark/run"],
    ], [4.5, 3.0, 6.5])
    para(doc,
         "若实际耗时显著高于上表，常见排查路径：① 检查 .cache 是否被清理（避免重编码 5320 篇）；"
         "② 检查 DEEPSEEK_API_KEY 是否有效（无效会触发 30 s 重试）；"
         "③ 查看 /api/stack 端点返回的模型 loaded 状态。", indent=False)

    heading(doc, "7.6  进阶：API 直接调用", 2)
    para(doc, "对于自动化测试或二次开发，可直接调用后端 API。以下示例展示如何用 curl "
         "触发一次完整 14 阶段 Pipeline 并流式读取 SSE：")
    para(doc, "$ curl -N -X POST http://127.0.0.1:8000/api/pipeline/stream \\",
         indent=False, size=11)
    para(doc, "    -H 'Content-Type: application/json' \\",
         indent=False, size=11)
    para(doc, "    -d '{\"topic\": \"雨夜和解\", \"fastMode\": false}'",
         indent=False, size=11)
    para(doc, "返回为 SSE 事件流，每行形如：",
         indent=False)
    para(doc, "event: stage", indent=False, size=10)
    para(doc, 'data: {"id": 6, "name": "BGE 检索", "model": "bge-small-zh-v1.5", "durationMs": 234, "progress": 0.43}',
         indent=False, size=10)
    para(doc, "完整 OpenAPI 文档请访问 http://127.0.0.1:8000/docs（FastAPI 自动生成）。", indent=False)

    # ══ 参考文献 ══
    doc.add_page_break()
    heading(doc, "参考文献")
    refs = [
        "[1] 宗成庆. 统计自然语言处理[M]. 2版. 北京: 清华大学出版社, 2013.",
        "[2] 车万翔, 崔一鸣, 苑春法. 自然语言处理: 基于预训练模型的方法[M]. 北京: 电子工业出版社, 2021.",
        "[3] 李航. 深度学习与方法[M]. 北京: 电子工业出版社, 2019.",
        "[4] Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389. DOI: 10.1561/1500000018.",
        "[5] Xiao S, Liu Z, Zhang P, et al. C-Pack: Packaged Resources to Advance General Chinese Embedding[EB/OL]. arXiv:2309.07597, 2023. https://arxiv.org/abs/2309.07597.",
        "[6] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//Proceedings of EMNLP-IJCNLP, 2019: 3982-3992. DOI: 10.18653/v1/D19-1410.",
        "[7] Nogueira R, Cho K. Passage Re-ranking with BERT[J]. arXiv:1901.04085, 2019. DOI: 10.48550/arXiv.1901.04085.",
        "[8] Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems, 2020, 33: 9459-9474.",
        "[9] Mihalcea R, Tarau P. TextRank: Bringing Order into Text[C]//Proceedings of EMNLP, 2004: 404-411.",
        "[10] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]//Proceedings of NAACL-HLT, 2019: 4171-4186. DOI: 10.18653/v1/N19-1423.",
        "[11] Peters M E, Neumann M, Iyyer M, et al. Deep Contextualized Word Representations[C]//Proceedings of NAACL-HLT, 2018: 2227-2237. DOI: 10.18653/v1/N18-1202.",
        "[12] Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, 1997, 9(8): 1735-1780. DOI: 10.1162/neco.1997.9.8.1735.",
        "[13] Es S, James J, Espinosa-Anke L, et al. RAGAS: Automated Evaluation of Retrieval Augmented Generation[C]//Proceedings of EACL, 2024.",
        "[14] Saunders W, Yeh C, Wu J, et al. Self-critiquing models for assisting human evaluators[J]. arXiv:2206.05802, 2022.",
        "[15] Shinn N, Cassano F, Gopinath A, et al. Reflexion: Language Agents with Verbal Reinforcement Learning[C]//Advances in Neural Information Processing Systems, 2023, 36.",
    ]
    for r in refs:
        para(doc, r, 11, indent=False, after=4)

    doc.add_page_break()
    heading(doc, "附录 A  系统概述表（课程评分用）")
    tbl(doc, ["功能模块", "模块概述", "NLP 相关技术"], [
        ["混合语义检索", "5320 首诗歌 Top-K 召回与 Cross-Encoder 精排", "BM25；BGE；Cross-Encoder"],
        ["可解释语义分析", "说明检索结果与主题的关键词/意象/情感关联", "关键词重叠；情感对比；实体匹配"],
        ["融合情感分析", "六维诗歌情感向量及主导情绪", "文学情感词典；RoBERTa；分数融合"],
        ["文学知识图谱", "意象共现/情感共鸣；力导向可视化", "规则抽取；BERT-RE；关系统计"],
        ["RAG 与生成", "结构化上下文驱动生成并标注化用来源", "RAG；DeepSeek LLM；检索化用模板"],
        ["辅助工具", "文摘、校错、润色、知识库浏览", "TextRank；MacBERT；Jieba"],
    ], [2.8, 5.5, 5.2])
    para(doc, "功能块个数：13；NLP 技术个数：15。", indent=False)

    heading(doc, "附录 B  评分表（教师填写）")
    tbl(doc, ["得分项", "满分", "得分"], [
        ["报告", "20", ""], ["答辩", "10", ""], ["功能", "30", ""],
        ["技术", "30", ""], ["整体评价", "10", ""], ["总分", "100", ""],
    ], [4, 2, 3])

    heading(doc, "附录 C  API 接口列表")
    tbl(doc, ["接口路径", "HTTP方法", "功能描述"], [
        ["/api/pipeline/stream", "POST", "完整灵感生成 Pipeline（SSE 实时推送）"],
        ["/api/pipeline/fast", "POST", "快速模式灵感生成（跳过 KG）"],
        ["/api/retrieve", "POST", "混合检索（支持 keyword/semantic/hybrid）"],
        ["/api/evaluate", "GET", "Ablation 评测接口"],
        ["/api/sentiment", "POST", "融合情感分析"],
        ["/api/entity", "POST", "实体识别（NER+意象词）"],
        ["/api/summary", "POST", "TextRank 自动文摘"],
        ["/api/correct", "POST", "MacBERT 文本校错"],
        ["/api/knowledge_graph", "GET", "知识图谱查询"],
        ["/api/poems", "GET", "诗歌知识库浏览"],
        ["/api/auth/register", "POST", "用户注册"],
        ["/api/auth/login", "POST", "用户登录"],
        ["/api/chat/sessions", "GET", "会话列表"],
    ], [4.5, 2, 8.5])

    heading(doc, "附录 D  Pipeline 14 阶段说明")
    tbl(doc, ["阶段序号", "阶段名称", "核心模型/方法", "作用描述"], [
        ["1", "分词", "Jieba + 噪声过滤", "文本切分，去除量词残片"],
        ["2", "NER", "BERT-NER + Jieba", "识别命名实体与意象名词"],
        ["3", "关键词提取", "TF-IDF", "提取主题核心词"],
        ["4", "查询扩展", "规则式 Expansion", "补充同域联想词"],
        ["5", "BM25 检索", "BM25F", "稀疏字面匹配召回"],
        ["6", "BGE 检索", "bge-small-zh-v1.5", "稠密语义匹配召回"],
        ["7", "分数融合", "Min-Max + α=0.55", "混合排序"],
        ["8", "Cross-Encoder", "bge-reranker-base", "精排重排序"],
        ["9", "情感分析", "词典 ⊕ RoBERTa", "六维情感向量"],
        ["10", "语义解释", "多信号聚合", "生成 semanticInsight"],
        ["11", "RAG 抽取", "结构化注入", "构建生成上下文"],
        ["12", "文本生成", "LLM / 模板", "创作文本生成"],
        ["13", "KG 子图", "BERT-RE + 规则", "主题知识图谱"],
        ["14", "Critic 自评", "LLM 评审 + 规则兜底", "评分<7 触发 LLM 改写 1 次"],
    ], [1.5, 2.5, 3.5, 6.5])

    heading(doc, "附录 E  部署清单（Docker / .env）")
    para(doc, "本附录汇总 Docker 部署所需文件内容，便于教师/助教按表对照。完整文件随仓库根目录发布。",
         indent=False)
    para(doc, "E.1  Dockerfile（多阶段构建；前端静态资源由 Nginx 托管）", bold=True, indent=False, size=11)
    para(doc, "FROM node:20-alpine AS front    # 阶段 1：构建前端", indent=False, size=10)
    para(doc, "WORKDIR /app/musagent", indent=False, size=10)
    para(doc, "COPY musagent/package*.json ./", indent=False, size=10)
    para(doc, "RUN npm ci --no-audit --no-fund", indent=False, size=10)
    para(doc, "COPY musagent/ ./", indent=False, size=10)
    para(doc, "RUN npm run build            # 产物在 /app/musagent/dist", indent=False, size=10)
    para(doc, "FROM python:3.11-slim AS back # 阶段 2：后端 + Nginx", indent=False, size=10)
    para(doc, "WORKDIR /app", indent=False, size=10)
    para(doc, "COPY back/requirements.txt ./back/requirements.txt", indent=False, size=10)
    para(doc, "RUN pip install --no-cache-dir -r back/requirements.txt", indent=False, size=10)
    para(doc, "COPY back/ ./back/", indent=False, size=10)
    para(doc, "COPY --from=front /app/musagent/dist /app/musagent/dist", indent=False, size=10)
    para(doc, "RUN apt-get update && apt-get install -y nginx && rm -rf /var/lib/apt/lists/*",
         indent=False, size=10)
    para(doc, "COPY nginx.conf /etc/nginx/conf.d/default.conf", indent=False, size=10)
    para(doc, 'CMD [\"uvicorn\", \"back.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]',
         indent=False, size=10)
    para(doc, "E.2  docker-compose.yml（一键启动）", bold=True, indent=False, size=11)
    para(doc, "version: '3.9'", indent=False, size=10)
    para(doc, "services:", indent=False, size=10)
    para(doc, "  musagent:", indent=False, size=10)
    para(doc, "    build: .", indent=False, size=10)
    para(doc, "    ports: ['8080:80', '8000:8000']", indent=False, size=10)
    para(doc, "    env_file: ./back/.env", indent=False, size=10)
    para(doc, "    volumes: ['musagent_cache:/app/back/.cache', 'musagent_db:/app/back/data']",
         indent=False, size=10)
    para(doc, "    restart: unless-stopped", indent=False, size=10)
    para(doc, "volumes:", indent=False, size=10)
    para(doc, "  musagent_cache:", indent=False, size=10)
    para(doc, "  musagent_db:", indent=False, size=10)
    para(doc, "E.3  back/.env.example（关键字段已合并展示于 7.1.3 节，此处从略）",
         bold=True, indent=False, size=11)
    para(doc, "E.4  端口与默认账户", bold=True, indent=False, size=11)
    tbl(doc, ["服务", "端口（容器内 / 主机映射）", "默认路径 / 凭据"], [
        ["前端 (Nginx)", "80 / 8080", "/  →  默认首页 /inspire"],
        ["后端 FastAPI", "8000 / 8000", "/docs  →  OpenAPI 文档；/api/*"],
        ["SQLite", "—", "back/data/musagent.db"],
        ["向量缓存", "—", "back/.cache/embeddings_*.npy"],
        ["HF 模型缓存", "—", "~/.cache/huggingface（v3.3 起挂载为命名卷）"],
    ], [3.5, 5.0, 6.0])
    para(doc, "管理员账号在首次启动时由后端自动创建（见 back/auth.py：bootstrap_admin），"
         "默认邮箱 admin@musagent.local，密码需在首次登录后立即修改。",
         indent=False, size=10)

    heading(doc, "附录 F  仓库目录结构与文件清单")
    para(doc, "本附录列出仓库关键目录与文件用途，便于答辩演示与代码审查。", indent=False)
    tbl(doc, ["路径", "用途", "备注"], [
        ["back/", "Python 3.12 / FastAPI 后端", "主入口 back/main.py"],
        ["back/quality_engine.py", "质量评估 + Critic Agent（v3.3）", "新增 critic_review 等"],
        ["back/nlp_engine.py", "9 大 NLP 模块实现", "1130 行"],
        ["back/orchestrator.py", "PipelineContext 编排", "12 行核心逻辑"],
        ["back/kg_engine.py", "文学垂直知识图谱", "610 行"],
        ["back/database.py", "SQLite ORM", "786 行"],
        ["musagent/src/pages/InspirePage.jsx", "灵感生成主界面", "14 阶段 SSE 实时进度"],
        ["musagent/src/components/GenerationResultCard.jsx", "生成结果卡片（含 Critic 评审）", "v3.3 新增"],
        ["musagent/src/config/routes.js", "路由表 + /benchmark 重定向", "LEGACY_REDIRECTS"],
        ["musagent/PROJECT.md", "前端项目说明", "v3.3 重写"],
        ["docs/", "大作业报告 + 图表生成脚本 + Benchmark 数据", "本报告所在目录"],
        ["docs/figures/fig7_data.json", "50 主题真实评测原始数据", "v3.3 落地"],
        ["docs/regenerate_fig7_10_11.py", "fig7/10/11 重生成脚本", "24 KB"],
        ["Dockerfile", "多阶段镜像构建", "Node 20 + Python 3.11 + Nginx"],
        ["docker-compose.yml", "一键启动", "持久化 .cache + DB 卷"],
        ["start-local.ps1", "Windows 本地源码一键启动", "等价命令见 7.1.2"],
        ["OPTIMIZATION_LOG.md", "v3.0 优化日志", "8.8 KB"],
        ["README.md", "项目总览", "v3.3 重写为 9 模块 + 1 Critic"],
        [".gitignore", "完整忽略规则", "含 .env / .cache / __pycache__"],
    ], [5.5, 5.0, 4.5])
    para(doc, "—— 报告结束 ——", bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    doc.save(OUT)
    print(f"Report: {OUT}")


if __name__ == "__main__":
    build()

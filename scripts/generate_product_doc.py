#!/usr/bin/env python3
"""生成 MusAgent 产品与技术白皮书（Word + 架构图）。"""
from __future__ import annotations

import textwrap
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
IMAGES = DOCS / "diagrams"
OUTPUT = DOCS / "MusAgent_产品文档.docx"

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

GOLD = "#e7d393"
BG = "#0f0f12"
CARD = "#1a1a22"
MUTED = "#888899"
BLUE = "#7ec8e3"
PURPLE = "#c9a0dc"
GREEN = "#98d8c8"


def _save(fig, name: str) -> Path:
    IMAGES.mkdir(parents=True, exist_ok=True)
    path = IMAGES / name
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor=BG)
    plt.close(fig)
    return path


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, indent=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    set_run_font(p.add_run(text), size=size, bold=bold)
    return p


def add_table(doc, headers, rows, widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            set_run_font(p.runs[0] if p.runs else p.add_run(h), size=9, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                if p.runs:
                    set_run_font(p.runs[0], size=9)
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def draw_tech_stack() -> Path:
    fig, ax = plt.subplots(figsize=(10, 5.2), facecolor=BG)
    ax.set_facecolor(BG)
    ax.axis("off")
    layers = [
        ("应用层", "React 19 · Vite 6 · D3 力导向图 · SSE 进度条", GOLD),
        ("编排层", "PipelineContext 13 阶段 · 耗时可观测 · 快速模式降级", BLUE),
        ("语义引擎", "Jieba · BM25 · BGE · Cross-Encoder · RoBERTa · BERT-NER/RE", PURPLE),
        ("知识层", "5320 首诗歌 · SQLite · 实体 4700+ · 关系 19000+", GREEN),
        ("生成层", "RAG 结构化注入 · DeepSeek Chat · 化用 citations", GOLD),
    ]
    y = 0.78
    for title, desc, color in layers:
        ax.add_patch(FancyBboxPatch((0.08, y - 0.1), 0.84, 0.14, boxstyle="round,pad=0.02",
                                    facecolor=CARD, edgecolor=color, linewidth=1.4))
        ax.text(0.12, y, title, color=color, fontsize=10, weight="bold", va="center")
        ax.text(0.28, y, desc, color="white", fontsize=8.5, va="center")
        y -= 0.16
    ax.text(0.5, 0.06, "Transformer-RAG v3 · 经典 NLP + 预训练模型融合 · 非纯 LLM 黑盒", ha="center", color=MUTED, fontsize=9)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    return _save(fig, "01_tech_stack.png")


def draw_pipeline_sse() -> Path:
    fig, ax = plt.subplots(figsize=(10, 4.8), facecolor=BG)
    ax.set_facecolor(BG)
    ax.axis("off")
    stages = [
        "分词", "NER", "查询扩展", "BM25+BGE", "Cross-Encoder",
        "情感融合", "RAG", "风格映射", "LLM", "质量评估", "KG子图",
    ]
    w, h = 0.075, 0.12
    xs = [0.04 + i * 0.086 for i in range(len(stages))]
    y = 0.58
    for i, (x, label) in enumerate(zip(xs, stages)):
        c = GOLD if label in ("LLM", "KG子图") else BLUE if "BM25" in label or "Cross" in label else PURPLE
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.01", facecolor=CARD, edgecolor=c, lw=1))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", color="white", fontsize=6.5)
        if i < len(stages) - 1:
            ax.annotate("", xy=(xs[i + 1], y + h / 2), xytext=(x + w, y + h / 2),
                        arrowprops=dict(arrowstyle="->", color=MUTED, lw=0.9))
    ax.add_patch(FancyBboxPatch((0.04, 0.22), 0.92, 0.22, boxstyle="round,pad=0.02", facecolor="#12121a", edgecolor=GREEN, lw=1))
    ax.text(0.5, 0.36, "POST /api/pipeline/stream  →  SSE event:stage  →  前端实时进度条 + 阶段耗时",
            ha="center", color=GREEN, fontsize=9)
    ax.text(0.5, 0.28, "快速模式：BM25-only 检索 · 跳过 KG 子图 · 约节省 40% 等待时间",
            ha="center", color=MUTED, fontsize=8)
    ax.text(0.5, 0.08, "13 阶段可解释流水线 — 每步标注模型名与 durationMs", ha="center", color=MUTED, fontsize=9)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    return _save(fig, "02_pipeline_sse.png")


def draw_retrieval() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.2), facecolor=BG)
    ax.set_facecolor(BG)
    ax.axis("off")
    boxes = [
        ("Query + 扩展词", 0.06, 0.55, BLUE),
        ("BM25 稀疏\nk1=1.5 b=0.75", 0.28, 0.62, PURPLE),
        ("BGE 稠密\nbge-small-zh-v1.5", 0.28, 0.38, PURPLE),
        ("α=0.55 融合\nMin-Max 归一化", 0.52, 0.5, BLUE),
        ("Cross-Encoder\nbge-reranker-base", 0.72, 0.5, GOLD),
        ("Top-5 + matchedTerms\n+ rerankScore", 0.52, 0.15, GREEN),
    ]
    for label, x, y, c in boxes:
        ax.add_patch(FancyBboxPatch((x, y), 0.18, 0.18, boxstyle="round,pad=0.02", facecolor=CARD, edgecolor=c, lw=1.2))
        ax.text(x + 0.09, y + 0.09, label, ha="center", va="center", color="white", fontsize=7.5)
    for (x1, y1, x2, y2) in [(0.24, 0.64, 0.28, 0.71), (0.24, 0.64, 0.28, 0.47), (0.46, 0.71, 0.52, 0.59),
                              (0.46, 0.47, 0.52, 0.55), (0.7, 0.59, 0.72, 0.59), (0.61, 0.5, 0.61, 0.33)]:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", color=MUTED, lw=1))
    ax.text(0.5, 0.03, "s_h = α·norm(BM25) + (1−α)·norm(cos_sim)  →  三级混合检索", ha="center", color=MUTED, fontsize=9)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    return _save(fig, "03_hybrid_retrieval.png")


def draw_semantic_kg() -> Path:
    fig, ax = plt.subplots(figsize=(10, 4.5), facecolor=BG)
    ax.set_facecolor(BG)
    ax.axis("off")
    ax.add_patch(FancyBboxPatch((0.05, 0.52), 0.42, 0.38, boxstyle="round,pad=0.02", facecolor=CARD, edgecolor=GOLD, lw=1.5))
    ax.text(0.26, 0.82, "诗内语义抽取", ha="center", color=GOLD, fontsize=10, weight="bold")
    for i, t in enumerate(["意象共现", "唤起情感", "含意象", "比喻/象征"]):
        ax.text(0.08, 0.72 - i * 0.08, f"• {t}", color="white", fontsize=8)
    ax.add_patch(FancyBboxPatch((0.53, 0.52), 0.42, 0.38, boxstyle="round,pad=0.02", facecolor=CARD, edgecolor=BLUE, lw=1.5))
    ax.text(0.74, 0.82, "跨诗语料统计", ha="center", color=BLUE, fontsize=10, weight="bold")
    for i, t in enumerate(["情感共鸣", "主题呼应", "语义呼应", "意象→情感"]):
        ax.text(0.56, 0.72 - i * 0.08, f"• {t}", color="white", fontsize=8)
    ax.add_patch(FancyBboxPatch((0.2, 0.12), 0.6, 0.28, boxstyle="round,pad=0.02", facecolor="#12121a", edgecolor=MUTED, lw=1))
    ax.text(0.5, 0.28, "元数据边（作者/体裁）置信度 ≤0.62 · 展示占比 ≤20%", ha="center", color=MUTED, fontsize=8.5)
    ax.text(0.5, 0.18, "D3 力导向图 · 关系类型均衡采样 · 左图右栏仪表盘", ha="center", color=GREEN, fontsize=8.5)
    ax.text(0.5, 0.04, "文学垂直 KG：文本—意象—情感 语义网络（非书目目录）", ha="center", color=MUTED, fontsize=9)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    return _save(fig, "04_semantic_kg.png")


def draw_vs_llm() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.2), facecolor=BG)
    ax.set_facecolor(BG)
    ax.axis("off")
    rows = [
        ("知识来源", "训练语料黑盒", "5320 首本地诗歌 + RAG 注入"),
        ("检索链路", "无", "BM25 + BGE + Cross-Encoder 三级"),
        ("情感理解", "隐式", "词典 ⊕ RoBERTa 六维融合"),
        ("关系图谱", "无", "意象共现 / 情感共鸣 / 主题呼应"),
        ("可解释性", "不可见", "semanticInsight + 阶段耗时 SSE"),
        ("质量保障", "无", "多维打分 + 化用 citations"),
        ("长文本", "通用续写", "超长档 4096 tokens + 分节"),
    ]
    ax.text(0.22, 0.92, "普通 LLM", ha="center", color=MUTED, fontsize=11, weight="bold")
    ax.text(0.72, 0.92, "MusAgent", ha="center", color=GOLD, fontsize=11, weight="bold")
    for i, (dim, llm, mus) in enumerate(rows):
        y = 0.8 - i * 0.11
        ax.text(0.04, y, dim, color=PURPLE, fontsize=8.5, va="center")
        ax.add_patch(FancyBboxPatch((0.18, y - 0.04), 0.32, 0.08, boxstyle="round,pad=0.01", facecolor=CARD, edgecolor=MUTED, lw=0.8))
        ax.text(0.34, y, llm, ha="center", va="center", color=MUTED, fontsize=7.5)
        ax.add_patch(FancyBboxPatch((0.56, y - 0.04), 0.38, 0.08, boxstyle="round,pad=0.01", facecolor=CARD, edgecolor=GOLD, lw=1))
        ax.text(0.75, y, mus, ha="center", va="center", color=GOLD, fontsize=7.5)
    return _save(fig, "05_vs_llm.png")


def build_doc(images: list[tuple[str, Path, str]]) -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    title = doc.add_heading("MusAgent 文学灵感平台", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("产品与技术白皮书"), size=14, bold=True, color=RGBColor(0x66, 0x66, 0x77))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta.add_run(f"版本 3.2  |  更新：{date.today().isoformat()}  |  Transformer-RAG v3"), size=10, color=RGBColor(0x88, 0x88, 0x99))

    doc.add_heading("摘要", level=1)
    add_para(doc, textwrap.dedent("""
        MusAgent 不是「把主题丢给大模型」的聊天机器人，而是一套面向诗词创作的
        语义检索增强生成（RAG）系统。系统在 5320 首现代诗与古典诗词语料上，
        串联 Jieba 分词、BERT-NER、BM25+BGE+Cross-Encoder 三级混合检索、
        文学词典与 RoBERTa 融合情感分析、可解释语义分析、文学垂直知识图谱、
        以及 DeepSeek 结构化生成与化用标注，形成「理解—检索—解释—生成—引用」
        的完整证据链。前端通过 SSE 实时推送 13 阶段 Pipeline 进度，
        知识图谱以 D3 力导向图展示意象—情感语义关系。
    """).strip())

    doc.add_heading("一、核心技术栈", level=1)
    add_para(doc, "系统采用经典 NLP 与预训练 Transformer 的混合架构，各模块职责清晰、可独立评测：")
    add_table(doc, ["模块", "技术方案", "模型/算法", "输出"], [
        ["分词", "Jieba + 诗歌噪声过滤", "HMM 新词识别", "words[], freq"],
        ["实体识别", "Hybrid NER", "ckiplab/bert-base-chinese-ner + 意象词典", "person/imagery/location"],
        ["关键词", "TF-IDF + 原文 boost", "倒排 df 统计", "Top-K keyword + tfidf"],
        ["检索", "三级混合召回", "BM25 + bge-small-zh-v1.5 + bge-reranker-base", "Top-5 + matchedTerms"],
        ["情感", "双通道融合", "EMOTION_DICT + uer/roberta-finetuned-jd", "六维向量 + dominant"],
        ["关系抽取", "规则 + BERT-RE", "hfl/chinese-bert-wwm-ext（可微调）", "三元组 head-rel-tail"],
        ["生成", "RAG + LLM", "DeepSeek Chat / 模板兜底", "文本 + citations"],
        ["校错", "MacBERT", "pycorrector", "修正列表"],
    ], [2.2, 2.8, 3.8, 4.5])

    doc.add_heading("二、13 阶段 NLP Pipeline + SSE 实时进度", level=1)
    add_para(doc, textwrap.dedent("""
        灵感生成由 orchestrator.PipelineContext 编排 13 个阶段，每阶段记录
        id、name、model、durationMs。POST /api/pipeline/stream 通过 Server-Sent Events
        推送 stage 事件，前端展示实时进度条与各步耗时，答辩/演示时可直观说明
        「系统究竟在做什么」。
    """).strip())
    add_para(doc, "阶段顺序（依赖约束）：分词 → NER → 查询扩展 → 混合检索 → 融合情感 → TextRank 摘要 → 可解释分析 → RAG 结构化抽取 → 风格映射 → LLM/模板生成 → 质量评估 → 主题 KG 子图。")
    add_para(doc, "快速模式（fastMode）：检索降级为 BM25-only，跳过 KG 子图构建，典型可节省约 40% 等待时间，适合课堂演示与迭代调试。")

    doc.add_heading("三、三级混合检索（技术核心）", level=1)
    add_para(doc, "诗歌主题往往极短（2—6 字），单一检索通道难以兼顾字面匹配与隐喻语义。Hybrid 模式执行：")
    for item in [
        "BM25 稀疏召回：对诗名、作者、典故词敏感，输出 matchedTerms 供可解释层使用",
        "BGE 稠密召回：bge-small-zh-v1.5 + retrieval instruction，弥补 paraphrase 字面失配",
        "分数融合：s_h = α·norm(BM25) + (1−α)·norm(cos_sim)，默认 α=0.55",
        "Cross-Encoder 精排：bge-reranker-base 对 Top-20 候选做 query-document 全注意力交互",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_para(doc, "Ablation 实验表明：BM25-only 与 Hybrid 的 Top-5 候选重叠数仅 2—4/5，证明稠密通道提供了互补语义信号而非重复排序。")

    doc.add_heading("四、融合情感分析与可解释层", level=1)
    add_para(doc, textwrap.dedent("""
        诗歌情感非简单正负二分类，而是喜悦/悲伤/平静/孤独/怀旧/激昂六维审美体验。
        analyze_sentiment 并行执行词典通道（EMOTION_DICT + TOPIC_EMOTION_HINTS）
        与 RoBERTa 通道（极性映射至六维），按 0.5:0.5 融合。
        检索结果 context_docs 可反哺情感判断——参考诗普遍呈「孤独」基调时，
        用户主题的情感推断会适度加权，体现「检索反哺语篇理解」。
    """).strip())
    add_para(doc, "explain_retrieval_results 将向量距离转化为自然语言：关键词命中、意象呼应、情感一致、Cross-Encoder 精排等 reasonTags + 逐条 summary + 全局 semanticInsight。")

    doc.add_heading("五、文学垂直知识图谱（语义关系优先）", level=1)
    add_para(doc, textwrap.dedent("""
        知识图谱面向「文本—意象—情感」语义关联，而非书目元数据目录。
        kg_engine 从诗歌正文实时抽取意象词（文学意象词典 + TF-IDF），
        构建诗内关系（imagery_co_occurs 意象共现、evokes_emotion 唤起情感、
        contains_imagery 含意象）与跨诗统计关系（emotion_resonance 情感共鸣、
        theme_echo 主题呼应）。作者/体裁等元数据边置信度刻意压至 ≤0.62，
        展示时 filter_edges 限制元数据占比 ≤20%。
    """).strip())
    add_para(doc, "数据规模（典型部署）：实体 4700+、关系 19000+；前端 ForceGraphView 力导向布局，关系类型均衡采样避免单一类型占满视图。")

    doc.add_heading("六、RAG 结构化注入与化用标注", level=1)
    add_para(doc, "Top-3 检索诗经二次 NLP 分析，抽取 keywords、emotion、adaptablePhrase 等结构化字段注入 Prompt。生成后输出 citations，标明「化用自《标题》·作者」，避免 RAG 沦为形式拼接。")

    doc.add_heading("七、账号体系与数据持久化", level=1)
    add_para(doc, "SQLite 持久化诗歌、实体、关系、反馈、对话会话。JWT 注册/登录，clientId 绑定 user_id 实现跨设备对话同步。")

    doc.add_heading("八、性能与工程优化", level=1)
    add_table(doc, ["优化项", "方案", "效果"], [
        ["SSE 流式", "pipeline/stream 分阶段推送", "用户可见进度，降低焦虑感"],
        ["快速模式", "BM25 + 跳过 KG", "约 −40% 耗时"],
        ["RAG 去重", "参考诗批量分词缓存", "避免重复 segment"],
        ["风格映射并行", "art+mus 合并 sty 阶段", "减少串行等待"],
        ["向量索引缓存", ".cache/embeddings_*.npy", "避免每次启动重编码 5320 篇"],
        ["长文本", "max_tokens 4096 + 分节 Prompt", "散文/短篇超长档"],
    ], [2.5, 4.5, 5.5])

    for heading, path, caption in images:
        doc.add_heading(heading, level=1)
        doc.add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            cap.runs[0].italic = True
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x77)

    doc.add_heading("附录：关键 API 端点", level=1)
    apis = [
        ("POST /api/pipeline/stream", "SSE 实时 13 阶段进度"),
        ("POST /api/pipeline", "完整 Pipeline（同步）"),
        ("GET /api/knowledge-graph", "力导向图数据源（语义关系优先）"),
        ("POST /api/auth/register · login", "JWT 账号体系"),
        ("GET /api/evaluate", "定量评测 + BM25/Hybrid Ablation"),
        ("GET /api/stack", "预训练模型栈 loaded 状态"),
        ("POST /api/kg/train-re", "BERT-RE 关系抽取微调"),
    ]
    add_table(doc, ["端点", "说明"], [[a, b] for a, b in apis], [5.5, 8])

    doc.add_heading("附录：测试与可复现性", level=1)
    add_para(doc, "product_regression_tests.py 含 24 项自动化断言，覆盖分词、检索、Pipeline 结构、可解释性、化用标注、知识库稳定性等。执行 python product_regression_tests.py 应输出 total_failures=0。")

    doc.save(OUTPUT)
    print(f"已生成：{OUTPUT}")


def main():
    images = [
        ("图1  五层技术架构", draw_tech_stack(), "应用—编排—语义引擎—知识—生成，Transformer-RAG v3"),
        ("图2  Pipeline + SSE 实时推送", draw_pipeline_sse(), "13 阶段可观测；快速模式可精简"),
        ("图3  三级混合检索", draw_retrieval(), "BM25 + BGE + Cross-Encoder，α=0.55 融合"),
        ("图4  文学语义知识图谱", draw_semantic_kg(), "意象/情感关系为主，元数据边降权"),
        ("图5  MusAgent vs 普通 LLM", draw_vs_llm(), "七维差异化：可检索、可解释、可引用"),
    ]
    build_doc(images)


if __name__ == "__main__":
    main()
